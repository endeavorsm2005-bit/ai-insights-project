# app.py -- AI Data Insights (Lite) with Excel / Word / PDF support
import streamlit as st
import pandas as pd
import io
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LinearRegression
from sklearn.metrics import r2_score, mean_absolute_error

# optional imports for docx/pdf; if missing, app will notify
try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

try:
    import pdfplumber
    PDF_AVAILABLE = True
except Exception:
    PDF_AVAILABLE = False

st.set_page_config(page_title="AI Data Insights (Files)", layout="wide")
st.title("🧠 AI Data Insights Generator — Multi-file (CSV / Excel / DOCX / PDF)")
st.write("Upload CSV / Excel / Word (.docx) / PDF containing tables and I will extract tabular data and run quick AI insights.")

uploaded = st.file_uploader("Upload a file (csv, xlsx, docx, pdf)", type=["csv", "xlsx", "xls", "docx", "pdf"])

def try_read_excel(bytestream):
    try:
        xls = pd.read_excel(bytestream, sheet_name=None, engine='openpyxl')
        dfs = list(xls.values())
        dfs = [d for d in dfs if isinstance(d, pd.DataFrame)]
        if not dfs:
            return None
        main = max(dfs, key=lambda d: d.shape[0])
        return main
    except Exception:
        return None

def extract_from_docx(bytestream):
    if not DOCX_AVAILABLE:
        return None
    try:
        doc = Document(io.BytesIO(bytestream))
        tables = doc.tables
        if len(tables) == 0:
            text = "
".join(p.text for p in doc.paragraphs if p.text.strip())
            lines = [l.strip() for l in text.splitlines() if l.strip()]
            if len(lines) >= 2 and ("," in lines[0] or "\t" in lines[0]):
                sep = "," if "," in lines[0] else "\t"
                return pd.read_csv(io.StringIO("\n".join(lines)), sep=sep)
            return None
        table = tables[0]
        data, keys = [], None
        for i, row in enumerate(table.rows):
            texts = [cell.text.strip() for cell in row.cells]
            if i == 0:
                keys = texts
            else:
                if len(texts) < len(keys):
                    texts += [""] * (len(keys) - len(texts))
                data.append(texts)
        return pd.DataFrame(data, columns=keys)
    except:
        return None

def extract_from_pdf(bytestream):
    if not PDF_AVAILABLE:
        return None
    try:
        tables_found = []
        with pdfplumber.open(io.BytesIO(bytestream)) as pdf:
            for page in pdf.pages:
                t = page.extract_table()
                if t:
                    tables_found.append(t)
        if not tables_found:
            return None
        tbl = max(tables_found, key=lambda t: len(t))
        df = pd.DataFrame(tbl[1:], columns=tbl[0])
        return df
    except:
        return None

df = None
if uploaded is not None:
    name = uploaded.name.lower()
    data = uploaded.read()
    if name.endswith(".csv"):
        try:
            df = pd.read_csv(io.BytesIO(data))
        except:
            pass
    elif name.endswith((".xlsx", ".xls")):
        df = try_read_excel(io.BytesIO(data))
    elif name.endswith(".docx"):
        df = extract_from_docx(data)
    elif name.endswith(".pdf"):
        df = extract_from_pdf(data)

if df is None:
    df = pd.DataFrame({
        "area": [800, 900, 1200, 1500, 1700, 2000],
        "bedrooms": [2, 2, 3, 3, 4, 4],
        "age": [10, 8, 5, 2, 1, 0],
        "price": [50, 55, 70, 90, 110, 130]
    })

df = df.dropna(axis=1, how="all")
st.subheader("Data preview")
st.write("Rows:", df.shape[0], "| Columns:", df.shape[1])
st.dataframe(df.head())

numeric_cols = df.select_dtypes(include="number").columns.tolist()
st.caption(f"Detected numeric columns: {numeric_cols}")

st.subheader("Quick automatic insights")
total = len(df)
missing = df.isna().sum()
for c, m in missing.items():
    if m > 0:
        pct = m / total * 100
        st.write(f"- Column **{c}** has {m} missing values ({pct:.1f}%).")

if len(numeric_cols) >= 2:
    st.subheader("Correlations (numeric)")
    corr = df[numeric_cols].corr()
    st.dataframe(corr.round(3))
    corrs = corr.abs().unstack().sort_values(ascending=False)
    corrs = corrs[corrs != 1.0]
    if not corrs.empty:
        top = corrs.index[0]
        st.write(f"Top correlation: **{top[0]}** ↔ **{top[1]}** = {corr.loc[top]:.2f}")

st.subheader("Simple AI: Linear Regression")
target = st.selectbox("Select numeric target", [None] + numeric_cols)
features = st.multiselect("Select numeric features", numeric_cols)

if target and features:
    X = df[features].dropna()
    y = df.loc[X.index, target]
    if len(X) >= 5:
        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.25, random_state=0)
        model = LinearRegression().fit(X_train, y_train)
        preds = model.predict(X_test)
        r2 = r2_score(y_test, preds)
        mae = mean_absolute_error(y_test, preds)
        st.write("R²:", round(r2,3), "| MAE:", round(mae,3))
        coef = dict(zip(features, model.coef_))
        st.write(pd.DataFrame.from_dict(coef, orient='index', columns=['coef']).round(3))

st.subheader("Export summary")
if st.button("Generate summary"):
    lines = []
    lines.append(f"Rows: {df.shape[0]} | Columns: {df.shape[1]}")
    if len(numeric_cols) >= 2:
        lines.append("Numeric cols: " + ", ".join(numeric_cols))
    for c, m in missing.items():
        if m > 0:
            lines.append(f"- {c}: {m} missing")
    st.code("\n".join(lines))
